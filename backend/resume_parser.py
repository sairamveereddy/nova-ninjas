"""
Resume Parser - Extracts text from PDF and DOCX files
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


async def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file content
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    try:
        import PyPDF2
        
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            logger.warning("PDF extraction returned empty text - might be image-based PDF")
            return ""
            
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file content
    
    Args:
        file_content: Raw bytes of the DOCX file
        
    Returns:
        Extracted text from the DOCX
    """
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


async def parse_resume(file_content: bytes, filename: str) -> str:
    """
    Parse resume file and extract text based on file type
    
    Args:
        file_content: Raw bytes of the file
        filename: Original filename to determine type
        
    Returns:
        Extracted text from the resume
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return await extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return await extract_text_from_docx(file_content)
    elif filename_lower.endswith('.doc'):
        raise ValueError("Old .doc format not supported. Please convert to .docx or .pdf")
    elif filename_lower.endswith('.txt'):
        return file_content.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type. Please upload PDF, DOCX, or TXT file")



def validate_magic_numbers(file_content: bytes, filename: str) -> bool:
    """
    Validate file content against magic numbers for its extension.
    
    Args:
        file_content: Raw bytes of the file
        filename: The filename to check extension against
        
    Returns:
        True if valid, False otherwise
    """
    filename_lower = filename.lower()
    
    # PDF Magic Number: %PDF- (25 50 44 46 2D)
    if filename_lower.endswith('.pdf'):
        return file_content.startswith(b'%PDF-')
        
    # DOCX Magic Number: PK (50 4B 03 04)
    elif filename_lower.endswith('.docx'):
        return file_content.startswith(b'PK\x03\x04')
        
    # TXT has no magic number, but we can check for binary content
    elif filename_lower.endswith('.txt'):
        try:
            file_content.decode('utf-8')
            return True
        except UnicodeDecodeError:
            return False
            
    return False


def validate_resume_file(filename: str, file_content: bytes) -> Optional[str]:
    """
    Validate resume file before processing
    
    Args:
        filename: The filename
        file_content: Raw bytes of the file
        
    Returns:
        Error message if invalid, None if valid
    """
    # Check file extension
    valid_extensions = ['.pdf', '.docx', '.txt']
    has_valid_ext = any(filename.lower().endswith(ext) for ext in valid_extensions)
    
    if not has_valid_ext:
        return "Invalid file type. Please upload PDF, DOCX, or TXT file"
    
    file_size = len(file_content)
    
    # Check file size (max 5MB - reduced from 10MB for security)
    max_size = 5 * 1024 * 1024  # 5MB
    if file_size > max_size:
        return "File too large. Maximum size is 5MB"
    
    # Check file size (min 10 bytes)
    min_size = 10
    if file_size < min_size:
        return f"Resume file looks empty ({file_size} bytes). Minimum required is {min_size} bytes."
        
    # Check magic numbers
    if not validate_magic_numbers(file_content, filename):
        return "Invalid file content. The file does not match its extension."
    
    return None


