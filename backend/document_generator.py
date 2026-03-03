"""
Document Generator - Creates optimized resumes and cover letters
"""

import os
import io
import logging
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aiohttp
import asyncio
import json
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, Union
from resume_analyzer import call_groq_api, unified_api_call, clean_json_response

logger = logging.getLogger(__name__)

# These are now imported from resume_analyzer.py
# GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
# GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
# GROQ_MODEL = "llama-3.3-70b-versatile"

# ============================================
# STRUCTURED RESUME SCHEMA (Step 3 & 4)
# ============================================

class ResumeHeader(BaseModel):
    full_name: str
    city_state: str
    phone: str
    email: str
    linkedin: str = ""
    portfolio: str = ""

class CoreSkills(BaseModel):
    languages: List[str] = []
    data_etl: List[str] = []
    cloud: List[str] = []
    databases: List[str] = []
    devops_tools: List[str] = []
    other: List[str] = []

class ExperienceRole(BaseModel):
    company: str
    job_title: str
    city_state_or_remote: str
    start: str
    end: str
    bullets: List[str]

class ProjectItem(BaseModel):
    name: str
    tech_stack: List[str] = []
    link: str = ""
    bullets: List[str]

class EducationItem(BaseModel):
    degree: Optional[str] = ""
    major: Optional[str] = ""
    university: Optional[str] = ""
    year: Optional[str] = ""

class ResumeDataSchema(BaseModel):
    """The structured data for a tailored resume"""
    header: ResumeHeader
    target_title: str
    positioning_statement: str
    core_skills: CoreSkills
    experience: List[ExperienceRole]
    projects: List[ProjectItem] = []
    education: List[EducationItem] = []
    certifications: List[str] = []

class ExpertTailoringOutput(BaseModel):
    """The complete response from the Expert AI"""
    alignment_highlights: List[str]
    cover_letter: str
    resume_data: ResumeDataSchema


def cleanup_bullet(s: str) -> str:
    """Removes weird bullet characters and extra whitespace (Step 4)"""
    import re
    cleaned = re.sub(r'[•●▪︎\-]', '', str(s))
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


def render_preview_text_from_json(data: Dict) -> str:
    """
    Converts the structured JSON from generate_optimized_resume_content 
    into a plain text format that ResumePaper can parse.
    """
    if not data:
        return ""
        
    out = []
    
    # Header
    contact = data.get("contactInfo", {})
    out.append(f"NAME\n{contact.get('name', 'Your Name')}")
    
    contacts = [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        contact.get("linkedin"),
        contact.get("website")
    ]
    contact_line = " | ".join([c for c in contacts if c])
    out.append(f"CONTACT\n{contact_line}")
    
    # Summary
    out.append(f"SUMMARY\n{data.get('summary', '')}")
    
    # Skills
    skills = data.get("skills", [])
    out.append(f"SKILLS\n" + "\n".join([f"• {s}" for s in skills]))
    
    # Experience
    out.append("EXPERIENCE")
    for job in data.get("experience", []):
        out.append(f"{job.get('company', '')} — {job.get('title', '')} | {job.get('location', '')}")
        out.append(f"{job.get('dates', '')}")
        for b in job.get("bullets", []):
            out.append(f"- {b}")
            
    # Projects
    if data.get("projects"):
        out.append("PROJECTS")
        for proj in data.get("projects", []):
            out.append(f"{proj.get('name', '')} — {proj.get('subtitle', '')}")
            for b in proj.get("bullets", []):
                out.append(f"- {b}")
                
    # Education
    if data.get("education"):
        out.append("EDUCATION")
        for edu in data.get("education", []):
            edu_line = f"{edu.get('degree', '')}, {edu.get('school', '')} | {edu.get('date', '')}"
            out.append(edu_line)
            
    return "\n\n".join(out)


def render_ats_resume_from_json(r: ResumeDataSchema) -> str:
    """
    Deterministic rendering of the ATS template (Step 4).
    Guarantees same headings, order, and style.
    """
    header_parts = [
        r.header.city_state,
        r.header.phone,
        r.header.email,
        r.header.linkedin,
        r.header.portfolio
    ]
    header_line = " | ".join([p for p in header_parts if p and str(p).strip() and str(p).lower() != "undefined"])
    import re
    header_line = re.sub(r'(?i)undefined', '', header_line)

    out = []
    # Force string and handle potential "undefined" or "None" leftovers
    fullName = str(r.header.full_name or "").strip()
    fullName = re.sub(r'(?i)^undefined\s*', '', fullName).strip()
    if not fullName or fullName.lower() in ["undefined", "none", "null"]:
        fullName = "Your Name"
    out.append(fullName)
    out.append(header_line)

    out.append("PROFESSIONAL SUMMARY")
    out.append(f"{r.target_title} — {r.positioning_statement}")

    out.append("SKILLS")
    if r.core_skills:
        if r.core_skills.languages: out.append(f"Languages: {', '.join(r.core_skills.languages)}")
        if r.core_skills.data_etl: out.append(f"Data/ETL: {', '.join(r.core_skills.data_etl)}")
        if r.core_skills.cloud: out.append(f"Cloud: {', '.join(r.core_skills.cloud)}")
        if r.core_skills.databases: out.append(f"Databases: {', '.join(r.core_skills.databases)}")
        if r.core_skills.devops_tools: out.append(f"DevOps/Tools: {', '.join(r.core_skills.devops_tools)}")
        if r.core_skills.other: out.append(f"Other: {', '.join(r.core_skills.other)}")

    if r.experience:
        out.append("EXPERIENCE")
        for role in r.experience:
            role_header = f"{role.company} — {role.job_title} | {role.city_state_or_remote}"
            out.append(role_header)
            out.append(f"{role.start} – {role.end}")
            for b in role.bullets:
                out.append(f"- {cleanup_bullet(b)}")

    if r.projects:
        out.append("PROJECTS")
        for p in r.projects:
            tech = ", ".join(p.tech_stack)
            link = f" ({p.link})" if p.link else ""
            out.append(f"{p.name} — {tech}{link}")
            for b in p.bullets:
                out.append(f"- {cleanup_bullet(b)}")

    if r.education:
        out.append("EDUCATION")
        for e in r.education:
            left_parts = [e.degree, e.major]
            left = ", ".join([p for p in left_parts if p and p.strip()])
            right_parts = [e.university, e.year]
            right = " | ".join([p for p in right_parts if p and p.strip()])
            out.append(f"{left} — {right}")

    if r.certifications:
        out.append("CERTIFICATIONS")
        for c in r.certifications:
            out.append(f"- {c}")

    # Final cleanup to ensure NO extra gaps
    return "\n".join([line for line in out if line.strip()]).strip()


async def generate_simple_tailored_resume(resume_text: str, job_description: str, job_title: str, company: str, byok_config: Optional[Dict] = None) -> str:
    """Integrated the user's Expert Resume Writer prompt for high-quality, compact output"""
    
    prompt = f"""
You are an expert resume writer and ATS optimization specialist.

TASK:
Convert the provided RAW_INPUT into a complete, ATS-friendly, recruiter-ready resume tailored for the JOB_DETAILS.
You must:
1) Extract and normalize all data (names, titles, dates, locations, bullets, skills, projects, education).
2) Fix formatting issues, duplicate sections, missing headers, inconsistent capitalization, and spacing.
3) Rewrite bullets to be strong, specific, and impact-focused. DO NOT invent facts or tools not present in the original for that time period.
4) Keep all original meaning. If a metric is unclear, keep it but don’t exaggerate.
5) Output in the requested RESUME_STYLE: ATS_ONE_PAGE_COMPACT (Ensure it is very dense and professional).

JOB_DETAILS:
Company: {company}
Title: {job_title}
Job Description: {job_description}

IMPORTANT RULES (non-negotiable):
- DO NOT add fake companies, degrees, awards, or certifications.
- DO NOT claim tools/tech that aren’t in the RAW_INPUT.
- DO NOT change job titles or dates unless clearly mis-ordered; if uncertain, keep as-is.
- DO NOT include “References available upon request.”
- Keep it ATS clean: no tables, no columns, no icons, no graphs, no fancy symbols.
- Use simple ASCII bullets like "-" only.
- Keep lines under ~110 characters where possible.
- If any section is missing a detail (ex: LinkedIn URL), keep the label but omit the value.

RESUME_STYLE: ATS_ONE_PAGE_COMPACT

OUTPUT FORMAT:
Return ONLY the resume text, nothing else.
Use exactly these section headers (omit any that truly have no content):
NAME
CONTACT
SUMMARY
SKILLS (Must be preserved and enhanced, never removed)
EXPERIENCE
PROJECTS
EDUCATION

- [STRICT HISTORICAL ACCURACY & ZERO HALLUCINATION] (MANDATORY):
    1. EXPERIENCE & PROJECTS: You MUST NOT add any technical skills, frameworks, tools, or software that are not explicitly mentioned in the [RAW_INPUT] for that specific role or project. 
    2. TIMELINE: Do not inject modern tech (like LLMs, Generative AI, RAG, Llama-3, etc.) into roles ending before 2022 unless they are explicitly in the source text.
    3. DATA INTEGRITY: Do not invent metrics or outcomes. If a JD requires a skill not in the resume, DO NOT fabricate it. Highlight a transferable existing skill instead.
- SUMMARY: 3–4 lines max, include your core stack + value.
- SKILLS: Grouped categories. Only include skills the candidate actually has in [RAW_INPUT].
- EXPERIENCE: For each role, 4–6 bullets max. Start each bullet with a strong verb. Use the EXACT dates provided.
- PROJECTS: 3 projects max. Preserve the exact tech stack mentioned in [RAW_INPUT].
- EDUCATION: Degree, school, location (if available), graduation year if present.

QUALITY BAR:
Your job is to make this resume look like a real senior engineer wrote it.
Make it tight, readable, and high-signal. Remove filler.

RAW_INPUT:
<<<
{resume_text}
>>>

Now generate the resume from RAW_INPUT:
"""
    try:
        logger.info(f"Running user-requested 'Expert' prompt tailoring for {company}")
        response = await unified_api_call(prompt, byok_config=byok_config, max_tokens=6000, model="llama-3.3-70b-versatile")
        
        if response and len(response.strip()) > 500:
            # Flatten to remove ALL excessive newlines
            import re
            content = response.strip()
            # Replace 3 or more newlines with 2
            content = re.sub(r'\n{3,}', '\n\n', content)
            return content
        
        # Immediate fallback to base resume if AI returns junk or empty
        return f"TAILORED RESUME: {job_title} at {company}\n\n" + resume_text
    except Exception as e:
        logger.error(f"Simple tailoring failed: {e}")
        return resume_text


async def generate_optimized_resume_content(resume_text: str, job_description: str, analysis: Dict, byok_config: Optional[Dict] = None, target_score: int = 85) -> Optional[Dict]:
    """Generate optimized resume content using AI - achieves target ATS score while preserving authenticity"""
    
    missing_skills = []
    if analysis.get("hardSkills", {}).get("missing"):
        missing_skills.extend([s.get("skill", s) if isinstance(s, dict) else s for s in analysis["hardSkills"]["missing"]])
    if analysis.get("softSkills", {}).get("missing"):
        missing_skills.extend([s.get("skill", s) if isinstance(s, dict) else s for s in analysis["softSkills"]["missing"]])
    
    keywords = analysis.get("keywordsToAdd", [])
    matched_skills = []
    if analysis.get("hardSkills", {}).get("matched"):
        matched_skills.extend([s.get("skill", s) if isinstance(s, dict) else s for s in analysis["hardSkills"]["matched"]])
    
    prompt = f"""
You are an Elite Resume Architect and ATS Optimization Expert. Your goal is to rewrite/enhance the resume to achieve a {target_score}% ATS match score.

=== ABSOLUTE RULES - VIOLATION = FAILURE ===
1. EVERY job position in the original MUST appear in output (count them!)
2. EVERY bullet point MUST be preserved - you can ONLY add more or enhance existing ones, never remove
3. EVERY project MUST be included with ALL its details
4. Output MUST be professional, high-signal, and achieve {target_score}% ATS compatibility
5. Copy the EXACT job titles, company names, dates - do NOT paraphrase
6. Keep the original professional summary structure but enhance it with keywords

=== YOUR TASK ===
Take each bullet point and ENHANCE it by:
- Naturally weaving in target keywords: {', '.join(missing_skills[:15] + keywords[:10])} ONLY IF RELEVANT to that specific role/time period.
- Adding market-standard industry keywords that recruiters and ATS systems look for.
- Adding specific metrics/numbers where logical (e.g., "improved efficiency by 20%").
- Using powerful action verbs (e.g., "orchestrated", "pioneered").
- WARNING: DO NOT add tools/technologies to a job if they were not released or not used in that era. DO NOT Hallucinate.
- Ensure the resulting content is dense with keywords relevant to: {job_description[:500]}

=== ORIGINAL RESUME (COPY ALL CONTENT - DO NOT SKIP ANY SECTION) ===
{resume_text}

=== TARGET JOB DESCRIPTION ===
{job_description}

=== OUTPUT FORMAT (JSON) ===
Return ONLY valid JSON, no markdown:

{{
    "contactInfo": {{
        "name": "COPY EXACT NAME",
        "email": "COPY EXACT EMAIL",
        "phone": "COPY EXACT PHONE",
        "location": "COPY EXACT LOCATION",
        "linkedin": "COPY IF EXISTS",
        "website": "COPY IF EXISTS"
    }},
    "summary": "ENHANCE the original summary to be 3-5 lines dense with high-impact keywords and your core value prop. Target {target_score}% alignment.",
    "experience": [
        {{
            "title": "COPY EXACT TITLE",
            "company": "COPY EXACT COMPANY NAME",
            "location": "COPY EXACT LOCATION",
            "dates": "COPY EXACT DATES",
            "bullets": [
                "ENHANCED version of original bullet 1 - integrate keywords and metrics",
                "ENHANCED version of original bullet 2",
                "ENHANCED version of original bullet 3",
                "... INCLUDE AND ENHANCE EVERY ORIGINAL BULLET ...",
                "ADD 1-2 extra high-impact bullets focusing on market-standard skills missed in original"
            ]
        }}
    ],
    "projects": [
        {{
            "name": "COPY EXACT PROJECT NAME",
            "subtitle": "Brief description or tech stack - enhance with keywords",
            "bullets": [
                "COPY and enhance all project details with impact metrics"
            ]
        }}
    ],
    "education": [
        {{
            "degree": "COPY EXACT DEGREE",
            "school": "COPY EXACT SCHOOL NAME",
            "location": "COPY IF EXISTS",
            "date": "COPY DATE"
        }}
    ],
    "skills": ["COPY ALL ORIGINAL SKILLS", "ADD market-standard technical and soft skills to reach {target_score}% match"]
}}

=== VERIFICATION CHECKLIST ===
- All positions included
- All bullets enhanced and preserved
- {target_score}% ATS optimization target met via keyword density
- Industry-standard terminology used throughout
- Metrics added to at least 50% of bullets

GENERATE THE {target_score}% OPTIMIZED RESUME JSON NOW:
"""

    try:
        # Use unified call with BYOK support
        response = await unified_api_call(prompt, byok_config=byok_config, max_tokens=8000, model="llama-3.1-8b-instant")
        if not response:
            return None
        
        # Robust cleaning using shared utility
        cleaned_json = clean_json_response(response)
        
        return json.loads(cleaned_json)
    except Exception as e:
        logger.error(f"Failed to generate resume content: {e}")
        return None


async def generate_cover_letter_content(resume_text: str, job_description: str, job_title: str, company: str, byok_config: Optional[Dict] = None) -> Optional[str]:
    """Generate a cover letter using AI"""
    
    prompt = f"""
Write a professional, compelling cover letter for this job application.

APPLICANT'S RESUME:
{resume_text[:2500]}

JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_description[:2000]}

Write a 3-4 paragraph cover letter that:
1. Opens with enthusiasm for the specific role and company
2. Highlights 2-3 most relevant experiences/achievements from the resume
3. Shows knowledge of the company/industry
4. Closes with a strong call to action

Use a professional but personable tone. Do NOT use brackets or placeholders.
Return ONLY the cover letter text, no JSON or markdown.
"""

    try:
        response = await unified_api_call(prompt, byok_config=byok_config)
        return response
    except Exception as e:
        logger.error(f"Failed to generate cover letter: {e}")
        return None


async def extract_compliance_facts(resume_text: str, byok_config: Optional[Dict] = None) -> Optional[Dict]:
    """Stage 1: Extract verbatim facts from resume into a strict JSON schema"""
    # Truncate resume text only if massive
    truncated_resume = resume_text[:15000]
    
    prompt = f"""
SYSTEM:
You are a precision fact extractor. Output JSON ONLY. No prose.

USER:
Extract EVERY detail verbatim from the candidate resume below. 

ABSOLUTE RULES:
1. DO NOT SKIP ANY JOBS. 
2. DO NOT SKIP ANY BULLET POINTS.
3. DO NOT SUMMARIZE. Copy text exactly as it appears.
4. Output ONLY valid JSON.

[CANDIDATE_BASE_RESUME]
<<<
{truncated_resume}
>>>

Return JSON with this exact schema:
{{
  "name": "",
  "location": "",
  "email": "",
  "phone": "",
  "links": {{"linkedin":"", "github":"", "portfolio":""}},
  "summary_original": "",
  "employers": [{{"company":"","title":"","location":"","start":"","end":"","bullets":[] }}],
  "projects": [{{"name":"","bullets":[],"tools":[],"metrics":[] }}],
  "education": [{{"degree":"","major":"","university":"","year":""}}],
  "skills": {{"technical":[], "soft":[], "certifications":[]}},
  "metrics_explicit": []
}}
"""
    try:
        # Use 70B for extraction to ensure NO content loss (Step 1 Optimization Reverted for Stability)
        response_text = await unified_api_call(prompt, byok_config=byok_config, max_tokens=3500, model="llama-3.3-70b-versatile")
        if not response_text:
            return {}
        
        json_text = clean_json_response(response_text)
        return json.loads(json_text)
    except Exception as e:
        logger.error(f"Fact extraction failed: {e}")
        return {}


async def generate_expert_documents(
    resume_text: str, 
    job_description: str, 
    user_info: Optional[Dict] = None, 
    byok_config: Optional[Dict] = None,
    selected_sections: Optional[List[str]] = None,
    selected_keywords: Optional[List[str]] = None
) -> Optional[Dict[str, Any]]:
    """Generate ATS Resume and Detailed CV using the compliance-grade two-stage pipeline"""
    
    # Resolving model selection (Optimization)
    extraction_model = "llama-3.1-8b-instant"
    drafting_model = "llama-3.3-70b-versatile"
    
    # Stage 1: Verbatim Fact Extraction
    logger.info("Stage 1: Extracting verbatim facts")
    facts_json = await extract_compliance_facts(resume_text, byok_config=byok_config)
    
    if not facts_json or not facts_json.get("employers"):
        logger.warning("Fact extraction flaked - rescuing with simple tailoring")
        simple_text = await generate_simple_tailored_resume(resume_text, job_description, "Target Role", "Target Company", byok_config=byok_config)
        cl_text = await generate_cover_letter_content(resume_text, job_description, "Target Role", "Target Company", byok_config=byok_config)
        return {
            "alignment_highlights": "- Full resume generated via rescue mode",
            "ats_resume": simple_text, "detailed_cv": simple_text,
            "cover_letter": cl_text, "resume_json": {}
        }

    # Resolve Header - Robust handling for "undefined" or "None" strings
    u = user_info or {}
    f = facts_json
    
    def get_clean_val(val, default):
        import re
        v = str(val or "").strip()
        v = re.sub(r'(?i)^undefined\s*', '', v).strip()
        if not v or v.lower() in ["undefined", "none", "null"]:
            return default
        return v

    header_info = {
        "full_name": get_clean_val(u.get("name") or f.get("name"), "Your Name"),
        "city_state": get_clean_val(u.get("location") or f.get("location"), "Location"),
        "phone": get_clean_val(u.get("phone") or f.get("phone"), "Phone Number"),
        "email": get_clean_val(u.get("email") or f.get("email"), "Email Address"),
        "linkedin": f.get("links", {}).get("linkedin", ""),
        "portfolio": f.get("links", {}).get("portfolio", "")
    }

    missing_skills_str = json.dumps(selected_keywords or [])
    selected_sections_str = json.dumps(selected_sections or [])

    resume_prompt = f"""
SYSTEM:
You are a resume tailoring engine. You will be given a resume and a job description.
Your ONLY goal is to make this resume competitive for THIS specific job.

═══════════════════════════════════════════
⚠️ GROUNDING CONSTRAINT — READ THIS FIRST
THIS IS THE MOST IMPORTANT INSTRUCTION
═══════════════════════════════════════════

You are an EDITOR, not a writer. You are NOT creating a resume. 
You are EDITING an existing resume that already exists below.

THE FOLLOWING ARE LOCKED AND CANNOT CHANGE UNDER ANY CIRCUMSTANCES:
- Candidate full name: must be copied EXACTLY from [RESUME_JSON]
- Email and phone: must be copied EXACTLY from [RESUME_JSON]  
- Every company name: must be copied EXACTLY from [RESUME_JSON]
- Every job title: must be copied EXACTLY from [RESUME_JSON]
- Every date range: must be copied EXACTLY from [RESUME_JSON]
- Every location: must be copied EXACTLY from [RESUME_JSON]
- Number of jobs: you CANNOT add or remove any job
- Number of projects: you CANNOT add or remove any project
- All metrics/numbers: must be copied EXACTLY from [RESUME_JSON]
- Education entries: must be copied EXACTLY from [RESUME_JSON]

IF A COMPANY, TITLE, OR NAME DOES NOT EXIST IN [RESUME_JSON], 
YOU ARE FORBIDDEN FROM USING IT. NO EXCEPTIONS.

BEFORE YOU WRITE A SINGLE WORD, verify:
- What is the candidate's real name in [RESUME_JSON]? → Use only that
- How many jobs are in [RESUME_JSON]? → Output must have exactly that many
- What are the exact company names? → Copy them character by character

VIOLATION TEST: After generating output, check every company name, 
every job title, and the candidate name against [RESUME_JSON]. 
If ANY of them don't match exactly, you have failed. Start over.

CRITICAL: Read the job description COMPLETELY before touching a single word of the resume.
CRITICAL: Every single change must be traceable to a specific requirement in the job description.
CRITICAL: Do NOT just reformat or relabel. Actually rewrite content to mirror JD language.

═══════════════════════════════════════════
YOUR INPUTS
═══════════════════════════════════════════

- [RESUME_JSON]: The candidate's full resume content
{json.dumps(facts_json, indent=2)}

- [JD_TEXT]: The job description
{job_description}

- [MISSING_SKILLS]: Array of skills the user selected to inject
{missing_skills_str}

- [SELECTED_SECTIONS]: Array of sections the user chose to enhance (if empty, assume ALL sections)
{selected_sections_str if selected_sections_str and selected_sections_str != "[]" else '["summary", "skills", "work_experience", "projects"]'}

═══════════════════════════════════════════
STEP 0 — READ AND EXTRACT FROM JD FIRST
═══════════════════════════════════════════

Before writing anything, extract and internally note:

A) JOB_TITLE: What is the exact role title?
B) TOP_5_KEYWORDS: What are the 5 most repeated or emphasized technical terms in the JD?
C) CORE_MISSION: What is the #1 problem this company is trying to solve with this hire?
D) REQUIRED_SKILLS: List every technical skill, framework, or methodology mentioned
E) PERSONA: What kind of engineer are they describing? (researcher? builder? both?)
F) MISSING_SKILLS_TO_ADD: {missing_skills_str}
G) SELECTED_SECTIONS: {selected_sections_str}

You must complete this extraction before proceeding to any section.

═══════════════════════════════════════════
STEP 1 — REWRITE SUMMARY
(only if "summary" is in SELECTED_SECTIONS)
═══════════════════════════════════════════

Rewrite the summary following this exact structure:

LINE 1 — Identity statement using JOB_TITLE language from the JD + candidate's years of experience
LINE 2 — Most relevant technical strength mapped to JD's CORE_MISSION  
LINE 3 — Second strongest alignment point with specific tools/methods from the JD
LINE 4 — (optional) Differentiator or scope statement

RULES:
✅ Use the exact terminology from the JD, not generic AI buzzwords
✅ If JD says "multi-agent systems" use "multi-agent systems" not "agentic workflows"
✅ If JD says "evaluation pipelines and guardrails" use that exact phrase
✅ Reference MISSING_SKILLS_TO_ADD naturally if candidate has adjacent experience
❌ Do NOT use phrases like "Proven track record" or "Delivering scalable solutions"
❌ Do NOT make the summary longer than 4 lines
❌ Do NOT invent experience that doesn't exist in the resume body

═══════════════════════════════════════════
STEP 2 — REBUILD SKILLS SECTION
(only if "skills" is in SELECTED_SECTIONS)
═══════════════════════════════════════════

Reorganize skills so the most JD-relevant categories appear FIRST.
Add ALL items from MISSING_SKILLS_TO_ADD to the appropriate category.
If a skill from the JD tech stack appears in MISSING_SKILLS_TO_ADD, it goes in skills.

Category order (most to least JD-relevant):
1. Agentic & Multi-Agent Systems (most JD-relevant category — list first)
2. LLM & GenAI
3. ML & Deep Learning  
4. Cloud & Infrastructure
5. Languages & APIs

RULES:
✅ Add every item from MISSING_SKILLS_TO_ADD
✅ Keep all existing skills — do not remove anything
✅ Group logically — don't dump everything in "Other"
❌ Do NOT create a category called "Other" — find the right home for every skill

═══════════════════════════════════════════
STEP 3 — TAILOR EXPERIENCE BULLETS
(only if "work_experience" is in SELECTED_SECTIONS)
═══════════════════════════════════════════

For EACH role, apply these 4 operations IN ORDER:

OPERATION A — REORDER
Move bullets that use JD keywords to the top.
Most JD-relevant bullet = position 1.

OPERATION B — SURFACE HIDDEN KEYWORDS
Look for bullets that describe JD concepts but use different words.
Update ONLY the terminology — keep the action and outcome identical.

Mandatory surfacing for this JD:
- Any bullet about "LangGraph/LangChain/CrewAI multi-step" 
  → add "multi-agent systems orchestration" to that bullet
- Any bullet about "citation verification" or "evidence gating" 
  → reframe as "guardrails for auditable, deterministic execution"  
- Any bullet about "evaluation harness" or "quality checks"
  → use "evaluation pipelines" (exact JD language)
- Any bullet about "RAG + retrieval"
  → connect to "reasoning and memory systems" (JD language)

OPERATION C — ADD IMPACT FRAMING
If a bullet has no business outcome, add one using the JD's domain.

OPERATION D — UPGRADE VERBS
Replace weak verbs with research-engineer level verbs matching JD persona:
"Collaborated on" → "Partnered to architect"
"Utilized" → "Engineered"  
"Developed" → "Designed and implemented"
"Enhanced" → "Optimized"

IF "work_experience" is NOT in SELECTED_SECTIONS:
- Only perform OPERATION A (reorder) — do not change any text

═══════════════════════════════════════════
STEP 4 — TAILOR PROJECTS
═══════════════════════════════════════════

ALWAYS tailor projects regardless of section selection.

Reorder projects so most JD-relevant appears first.
For this JD, correct order is:
1. JobNinjas AI Copilot (most relevant — agentic workflow, end-to-end builder)
2. Enterprise Patent Assistant (RAG + evaluation)
3. Text-to-SQL Copilot (guardrails + evaluation)

For each project bullet:
✅ Surface multi-agent / agentic language where it genuinely exists
✅ Keep ALL metrics exactly as written — never change numbers
✅ Connect project outcomes to JD pain points where honest

═══════════════════════════════════════════
ABSOLUTE RULES — NEVER VIOLATE THESE
═══════════════════════════════════════════

❌ NEVER leave summary blank or with a dash
❌ NEVER fabricate tools, companies, dates, or metrics
❌ NEVER add new bullet points — only edit and reorder existing ones  
❌ NEVER change company names, job titles, locations, or dates
❌ NEVER change any metric numbers
❌ NEVER add skills to experience bullets that have zero supporting evidence
❌ NEVER use filler phrases: "proven track record", "results-driven", "passionate about"
✅ ALWAYS return a complete resume — no section left empty or with placeholder
✅ ALWAYS use the JD's exact terminology when the candidate's experience supports it
✅ ALWAYS add every item from MISSING_SKILLS_TO_ADD to the skills section

═══════════════════════════════════════════
SELF-CHECK BEFORE RETURNING OUTPUT
═══════════════════════════════════════════

Before returning, verify:
[ ] Summary mentions the JD role title or close equivalent
[ ] Summary contains at least 2 phrases from the JD (not from the original resume)
[ ] Every item in MISSING_SKILLS_TO_ADD appears in the skills section
[ ] At least 3 experience bullets have been updated with JD terminology
[ ] No section is blank or contains only a dash
[ ] No metrics were changed
[ ] No new bullet points were added
[ ] Projects are reordered with most JD-relevant first

If ANY of these checks fail, fix before returning.

═══════════════════════════════════════════
OUTPUT FORMAT — STRICT JSON
═══════════════════════════════════════════

{{
  "jd_extraction": {{
    "job_title": "...",
    "top_5_keywords": ["...", "...", "...", "...", "..."],
    "core_mission": "...",
    "persona": "..."
  }},
  "tailored_resume": {{
    "name": "SAIRAM KUMAR",
    "contact": "...",
    "summary": ["line 1", "line 2", "line 3"],
    "skills": {{
      "Agentic & Multi-Agent Systems": ["..."],
      "LLM & GenAI": ["..."],
      "ML & Deep Learning": ["..."],
      "Cloud & Infrastructure": ["..."],
      "Languages & APIs": ["..."]
    }},
    "experience": [
      {{
        "company": "exact name",
        "title": "exact title",
        "location": "exact location",
        "dates": "exact dates",
        "bullets": ["bullet 1", "bullet 2", "..."]
      }}
    ],
    "projects": [
      {{
        "name": "exact name",
        "tech": "tech stack",
        "bullets": ["bullet 1", "bullet 2", "bullet 3"]
      }}
    ],
    "education": [
      {{"degree": "...", "school": "..."}}
    ]
  }},
  "cover_letter": "A brief, highly compelling 3-paragraph cover letter tailored to the job description and the candidate's core strengths.",
  "changes": [
    {{
      "section": "summary|skills|experience|projects",
      "company_or_project": "which role/project this applies to",
      "type": "rewritten|keyword_surfaced|reordered|verb_upgraded|skill_added|impact_added",
      "original": "exact original text",
      "updated": "exact new text",
      "jd_reason": "specific JD phrase or requirement this serves"
    }}
  ],
  "skills_added": ["list of MISSING_SKILLS successfully injected"],
  "self_check_passed": true
  "skills_added": ["list of MISSING_SKILLS successfully injected"],
  "self_check_passed": true
}}

OUTPUT ONLY VALID JSON. NO PREAMBLE. NO CHAT. NO CONVERSATIONAL TEXT.
"""

    try:
        response_text = await unified_api_call(resume_prompt, byok_config=byok_config, max_tokens=6000, model=drafting_model)
        if not response_text:
            raise ValueError("Drafting failed")
            
        json_text = clean_json_response(response_text)
        raw_output = json.loads(json_text)
        
        # Merge cover letter back into output for compatibility if model outputted top-level
        cl_text = raw_output.get('cover_letter', "")
        
        # Validate and Render
        class JDExtraction(BaseModel):
            job_title: Optional[str] = ""
            top_5_keywords: Optional[List[str]] = []
            core_mission: Optional[str] = ""
            persona: Optional[str] = ""
            
        class ResumeChangeItem(BaseModel):
            section: str
            company_or_project: Optional[str] = None
            type: Optional[str] = ""
            original: Optional[str] = None
            updated: str
            jd_reason: Optional[str] = ""

        class ExperienceNew(BaseModel):
            company: Optional[str] = ""
            title: Optional[str] = ""
            location: Optional[str] = ""
            dates: Optional[str] = ""
            start: Optional[str] = ""
            end: Optional[str] = ""
            bullets: List[str] = []

        class ProjectNew(BaseModel):
            name: Optional[str] = ""
            tech: Optional[str] = ""
            bullets: List[str] = []

        class EducationNew(BaseModel):
            degree: Optional[str] = ""
            school: Optional[str] = ""
            university: Optional[str] = ""
            major: Optional[str] = ""
            year: Optional[str] = ""

        class TailoredResume(BaseModel):
            name: Optional[str] = ""
            contact: Optional[Any] = ""
            location: Optional[str] = ""
            email: Optional[str] = ""
            phone: Optional[str] = ""
            links: Optional[Dict[str, str]] = {}
            summary: List[str] = []
            skills: Dict[str, List[str]] = {}
            experience: List[ExperienceNew] = []
            projects: List[ProjectNew] = []
            education: List[EducationNew] = []

        class ExpertResumeOutput(BaseModel):
            jd_extraction: Optional[JDExtraction] = None
            tailored_resume: TailoredResume
            cover_letter: Optional[str] = ""
            changes: List[ResumeChangeItem] = []
            skills_added: List[str] = []
            self_check_passed: Optional[bool] = True

        validated = ExpertResumeOutput(**raw_output)
        t_resume = validated.tailored_resume

        # Build backward-compatible ResumeDataSchema
        def safe_clean(val):
            if not val: return ""
            import re
            return re.sub(r'(?i)^undefined\s*', '', str(val)).strip()

        rd_header = {
            "full_name": safe_clean(t_resume.name) or get_clean_val(header_info.get("full_name"), "Your Name"),
            "city_state": safe_clean(t_resume.location) or get_clean_val(header_info.get("city_state"), ""),
            "phone": safe_clean(t_resume.phone) or get_clean_val(header_info.get("phone"), ""),
            "email": safe_clean(t_resume.email) or get_clean_val(header_info.get("email"), ""),
            "linkedin": get_clean_val(header_info.get("linkedin"), ""),
            "portfolio": get_clean_val(header_info.get("portfolio"), "")
        }

        if t_resume.links:
            rd_header["linkedin"] = safe_clean(t_resume.links.get("linkedin")) or rd_header["linkedin"]
            rd_header["portfolio"] = safe_clean(t_resume.links.get("github")) or safe_clean(t_resume.links.get("portfolio")) or rd_header["portfolio"]

        if isinstance(t_resume.contact, dict):
            rd_header["city_state"] = safe_clean(t_resume.contact.get("location")) or rd_header["city_state"]
            rd_header["email"] = safe_clean(t_resume.contact.get("email")) or rd_header["email"]
            rd_header["phone"] = safe_clean(t_resume.contact.get("phone")) or rd_header["phone"]
            if "links" in t_resume.contact:
                links = t_resume.contact["links"]
                rd_header["linkedin"] = safe_clean(links.get("linkedin")) or rd_header["linkedin"]
                rd_header["portfolio"] = safe_clean(links.get("github")) or safe_clean(links.get("portfolio")) or rd_header["portfolio"]

        derived_title = "Professional"
        if validated.jd_extraction and validated.jd_extraction.job_title:
            derived_title = validated.jd_extraction.job_title[:50]
        else:
            derived_title = u.get("target_role", "Professional")

        rd = ResumeDataSchema(
            header=rd_header,
            target_title=derived_title,
            positioning_statement="\\n".join(t_resume.summary),
            core_skills=CoreSkills(
                languages=t_resume.skills.get("Languages & APIs", []),
                data_etl=[],
                cloud=t_resume.skills.get("Cloud & Infrastructure", []),
                databases=[],
                devops_tools=t_resume.skills.get("Agentic & Multi-Agent Systems", []) + t_resume.skills.get("Agentic Frameworks", []),
                other=t_resume.skills.get("AI & LLM", []) + t_resume.skills.get("ML & Deep Learning", [])
            ),
            experience=[
                ExperienceRole(
                    company=e.company or "",
                    job_title=e.title or "",
                    city_state_or_remote=e.location or "",
                    start=e.start if e.start else (e.dates.split(" - ")[0] if e.dates and " - " in e.dates else e.dates),
                    end=e.end if e.end else (e.dates.split(" - ")[1] if e.dates and " - " in e.dates else "Present"),
                    bullets=e.bullets
                ) for e in t_resume.experience
            ],
            projects=[
                ProjectItem(
                    name=p.name or "",
                    tech_stack=p.tech.split(", ") if isinstance(p.tech, str) and p.tech else [],
                    link="",
                    bullets=p.bullets
                ) for p in t_resume.projects
            ],
            education=[
                EducationItem(
                    degree=ed.degree or "",
                    university=ed.university or ed.school or "",
                    major=ed.major or "",
                    year=ed.year or ""
                ) for ed in t_resume.education
            ],
            certifications=[]
        )

        ats_resume_text = render_ats_resume_from_json(rd)
        
        return {
            "alignment_highlights": "",
            "ats_resume": ats_resume_text, "detailed_cv": ats_resume_text,
            "cover_letter": cl_text, "resume_json": rd.model_dump(),
            "changes": [
                {
                    "section": c.section,
                    "type": c.type,
                    "original": c.original,
                    "updated": c.updated,
                    "reason": c.jd_reason
                } for c in validated.changes
            ],
            "skills_added": validated.skills_added,
            "skills_skipped": []
        }
    except Exception as e:
        logger.error(f"Stage 2 Expert failed: {e}")
        import traceback
        with open("debug_expert_error.txt", "w", encoding="utf-8") as tf:
            tf.write("ERROR:\\n" + traceback.format_exc() + "\\n\\n")
            tf.write("RAW JSON:\\n" + str(locals().get('json_text', 'No json_text')))
            
        simple_text = await generate_simple_tailored_resume(resume_text, job_description, "Position", "Company", byok_config=byok_config)
        cl_text = await generate_cover_letter_content(resume_text, job_description, "Position", "Company", byok_config=byok_config)
        return {
            "alignment_highlights": "- Tailored analysis complete",
            "ats_resume": simple_text, "detailed_cv": simple_text,
            "cover_letter": cl_text, "resume_json": {},
            "changes": [], "skills_added": [], "skills_skipped": []
        }

    return None

def create_resume_docx(resume_data: Dict, font_family: str = "Times New Roman") -> io.BytesIO:
    """Create a comprehensive Word document from resume data"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_family
    font.size = Pt(10)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Contact Info / Header
    contact = resume_data.get("contactInfo", {})
    name = contact.get("name", "Your Name")
    
    # Name - Large and Bold
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)
    
    # Contact details - First line (email, phone, location)
    contact_line1 = []
    if contact.get("email"):
        contact_line1.append(contact["email"])
    if contact.get("phone"):
        contact_line1.append(contact["phone"])
    if contact.get("location"):
        contact_line1.append(contact["location"])
    
    if contact_line1:
        contact_para1 = doc.add_paragraph(" | ".join(contact_line1))
        contact_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para1.paragraph_format.space_after = Pt(0)
    
    # Contact details - Second line (linkedin, website)
    contact_line2 = []
    if contact.get("linkedin"):
        contact_line2.append(contact["linkedin"])
    if contact.get("website"):
        contact_line2.append(contact["website"])
    
    if contact_line2:
        contact_para2 = doc.add_paragraph(" | ".join(contact_line2))
        contact_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para2.paragraph_format.space_after = Pt(2)
    
    # Professional Summary
    if resume_data.get("summary"):
        heading = doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        heading.runs[0].font.size = Pt(12)
        summary_para = doc.add_paragraph(resume_data["summary"])
        summary_para.paragraph_format.space_after = Pt(2)
    
    # Professional Experience
    if resume_data.get("experience"):
        heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for i, exp in enumerate(resume_data["experience"]):
            # Title and Company on same line
            exp_header = doc.add_paragraph()
            title_run = exp_header.add_run(f"{exp.get('title', 'Title')}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            company_text = f" | {exp.get('company', 'Company')}"
            if exp.get("location"):
                company_text += f", {exp['location']}"
            exp_header.add_run(company_text)
            exp_header.paragraph_format.space_after = Pt(0)
            
            # Dates
            if exp.get("dates"):
                dates_para = doc.add_paragraph()
                dates_run = dates_para.add_run(exp["dates"])
                dates_run.italic = True
                dates_para.paragraph_format.space_after = Pt(0)
            
            # Bullets - all of them
            bullets = exp.get("bullets", [])
            for bullet in bullets:
                if bullet and bullet.strip():
                    bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                    bullet_para.paragraph_format.space_after = Pt(0)
            
            # Add space between experiences
            if i < len(resume_data["experience"]) - 1:
                pass # Removed extra spacer between roles
    
    # Projects Section
    if resume_data.get("projects"):
        heading = doc.add_heading("PROJECTS", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for project in resume_data["projects"]:
            proj_header = doc.add_paragraph()
            name_run = proj_header.add_run(f"{project.get('name', 'Project')}")
            name_run.bold = True
            
            # Handle subtitle or technologies
            subtitle = project.get("subtitle") or project.get("technologies")
            if subtitle:
                proj_header.add_run(f" — {subtitle}")
            proj_header.paragraph_format.space_after = Pt(0)
            
            if project.get("description"):
                desc_para = doc.add_paragraph(project["description"])
                desc_para.paragraph_format.space_after = Pt(2)
            
            # Handle bullets - could be list or single string
            bullets = project.get("bullets", [])
            if isinstance(bullets, str):
                bullets = [bullets]
            for bullet in bullets:
                if bullet and bullet.strip():
                    # Check if it's an "Impact:" section
                    if bullet.lower().startswith("impact:"):
                        impact_para = doc.add_paragraph()
                        impact_label = impact_para.add_run("Impact: ")
                        impact_label.bold = True
                        impact_para.add_run(bullet[7:].strip())
                    else:
                        doc.add_paragraph(bullet, style='List Bullet')
    
    # Education
    if resume_data.get("education"):
        heading = doc.add_heading("EDUCATION", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for edu in resume_data["education"]:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(f"{edu.get('degree', 'Degree')}")
            degree_run.bold = True
            edu_para.add_run(f" | {edu.get('school', 'School')}")
            if edu.get("location"):
                edu_para.add_run(f", {edu['location']}")
            edu_para.paragraph_format.space_after = Pt(0)
            
            # Date and GPA
            details_parts = []
            if edu.get("date"):
                details_parts.append(edu["date"])
            if edu.get("gpa"):
                details_parts.append(f"GPA: {edu['gpa']}")
            if details_parts:
                details_para = doc.add_paragraph(" | ".join(details_parts))
                details_para.runs[0].italic = True
                details_para.paragraph_format.space_after = Pt(2)
            
            if edu.get("details"):
                doc.add_paragraph(edu["details"])
    
    # Skills - Handle both list and dict formats
    skills_data = resume_data.get("skills", {})
    if skills_data:
        heading = doc.add_heading("SKILLS", level=1)
        heading.runs[0].font.size = Pt(12)
        
        if isinstance(skills_data, dict):
            # New format with categories
            if skills_data.get("technical"):
                tech_para = doc.add_paragraph()
                tech_label = tech_para.add_run("Technical: ")
                tech_label.bold = True
                tech_para.add_run(" • ".join(skills_data["technical"]))
            
            if skills_data.get("tools"):
                tools_para = doc.add_paragraph()
                tools_label = tools_para.add_run("Tools & Platforms: ")
                tools_label.bold = True
                tools_para.add_run(" • ".join(skills_data["tools"]))
            
            if skills_data.get("other"):
                other_para = doc.add_paragraph()
                other_label = other_para.add_run("Other: ")
                other_label.bold = True
                other_para.add_run(" • ".join(skills_data["other"]))
        else:
            # Old format - simple list
            skills_text = " • ".join(skills_data)
            doc.add_paragraph(skills_text)
    
    # Certifications
    if resume_data.get("certifications"):
        heading = doc.add_heading("CERTIFICATIONS", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for cert in resume_data["certifications"]:
            cert_para = doc.add_paragraph()
            cert_name = cert_para.add_run(f"{cert.get('name', 'Certification')}")
            cert_name.bold = True
            if cert.get("issuer"):
                cert_para.add_run(f" - {cert['issuer']}")
            if cert.get("date"):
                cert_para.add_run(f" ({cert['date']})")
    
    # Additional Sections (Awards, Publications, Volunteer, etc.)
    if resume_data.get("additionalSections"):
        for section in resume_data["additionalSections"]:
            if section.get("title") and section.get("items"):
                heading = doc.add_heading(section["title"].upper(), level=1)
                heading.runs[0].font.size = Pt(12)
                
                for item in section["items"]:
                    if item and item.strip():
                        doc.add_paragraph(f"• {item}")
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def create_cover_letter_docx(cover_letter_text: str, job_title: str, company: str, font_family: str = "Times New Roman") -> io.BytesIO:
    """Create a Word document for the cover letter"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_family
    font.size = Pt(11)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Date
    from datetime import datetime
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Spacing
    
    # Greeting
    doc.add_paragraph(f"Re: Application for {job_title} at {company}")
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph()
    
    # Body - split into paragraphs
    paragraphs = cover_letter_text.strip().split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("[Your Name]")
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_text_docx(text: str, title: str = "Document", font_family: str = "Times New Roman", template: str = "standard") -> io.BytesIO:
    """Create a Word document from raw text, preserving basic formatting and styling"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_family
    font.size = Pt(10)

    is_modern = template == 'modern'
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Split text into lines and add to document
    lines = text.strip().split('\n')
    
    # First line is usually the name
    if lines:
        name_para = doc.add_paragraph()
        # Strip literal "NAME:" prefix if AI included it 
        name_text = lines[0].strip()
        if name_text.upper().startswith("NAME:"):
            name_text = name_text[5:].strip()
            
        name_run = name_para.add_run(name_text.upper())
        name_run.bold = True
        name_run.font.size = Pt(18)  # Reduced from 24pt for better appearance
        if not is_modern:
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)
        lines = lines[1:]

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
            
        if line_stripped.startswith('===') and line_stripped.endswith('==='):
            # Section header
            p = doc.add_paragraph()
            clean_title = line_stripped.replace('=', '').strip().upper()
            
            # Skip literal generic headers that AI might provide
            if clean_title in ["NAME", "CONTACT"]:
                continue
                
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(11)
            if not is_modern:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_bottom_border(p)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif (line_stripped.isupper() and len(line_stripped) < 50) or (line_stripped.startswith('#') and len(line_stripped) < 60) or (line_stripped.upper() in ["EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EDUCATION", "PROJECTS", "SKILLS", "SUMMARY", "PROFESSIONAL SUMMARY", "PROFILE", "CERTIFICATIONS", "AWARDS", "LANGUAGES", "TECHNICAL SKILLS", "CORE COMPETENCIES"]):
            # Main sections like PROFESSIONAL EXPERIENCE
            clean_title = line_stripped.replace('#', '').strip().upper()
            
            # Skip literal generic headers
            if clean_title in ["NAME", "CONTACT"]:
                continue
                
            p = doc.add_paragraph()
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(11)
            if not is_modern:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_bottom_border(p)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        elif line_stripped.startswith('- ') or line_stripped.startswith('• '):
            # Bullet point
            p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
        else:
            # Check if this is a job header (contains separators like " — " or " | ")
            is_job_header = " — " in line_stripped or " | " in line_stripped or " – " in line_stripped
            
            p = doc.add_paragraph()
            if is_job_header:
                run = p.add_run(line)
                run.bold = True
            else:
                # Try to center contact info if it's the first non-empty line after the name
                # Strip literal "CONTACT:" prefix if AI included it
                if i == 0 and line_stripped.upper().startswith("CONTACT:"):
                     p.clear()
                     p.add_run(line_stripped[8:].strip())
                     
                if i == 0 and ('@' in line or '|' in line) and not is_modern:
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.add_run(line)

            p.paragraph_format.space_after = Pt(0)
            
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
